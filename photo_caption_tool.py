import configparser
import csv
import math
import os
import platform
import re
import shutil
import subprocess
from pathlib import Path

from PIL import Image, ImageOps, ImageFont, ImageDraw
from pillow_heif import register_heif_opener
from docx import Document
from docx.shared import Inches, Mm


register_heif_opener()
configs = configparser.ConfigParser(comment_prefixes="|", allow_no_value=True)
configs.optionxform = str
images_directory = ""
all_images_exif_data = {}
rotation = ["1", "8", "3", "6"]  # Rotation of images, as represented in EXIF
valid_actions = []


class highlight:
    def red(thetext):
        return f"\033[31m{thetext}\033[0m"

    def green(thetext):
        return f"\033[32m{thetext}\033[0m"

    def bold(thetext):
        return f"\033[1m{thetext}\033[0m"


def _build_new_caption(project, site, subject, description) -> str:
    caption = ""
    if project:
        caption = f"Project: {project}. "
    if site:
        caption += f"Site: {site}. "
    if subject:
        caption += f"{subject}: "
    caption += description
    return caption


def _display_menu() -> str:
    global valid_actions
    if images_directory:
        print(f"\nPhotos Folder: {highlight.green(images_directory)}")
    else:
        print(f"\nPhotos Folder: {highlight.red('(not set)')}")
    print("----------------------------")
    print(" Choose an action:")
    print(f" {highlight.bold('1')} - Load Photos from Folder")
    valid_actions = ["1"]
    if images_directory:
        csv_file = Path(images_directory) / "Photo Log.csv"
        word_doc = Path(images_directory) / "Contact Sheet.docx"
        print(f" {highlight.bold('2')} - Create New Photo Log")
        valid_actions.append("2")
        if csv_file.is_file():
            print(f" {highlight.bold('3')} - View CSV File")
            valid_actions.append("3")
            print(f" {highlight.bold('4')} - Copy and Rename Photos")
            valid_actions.append("4")
            print(f" {highlight.bold('5')} - Annotate and Rename Photos")
            valid_actions.append("5")
            print(f" {highlight.bold('6')} - Create Contact Sheet")
            valid_actions.append("6")
            if word_doc.is_file():
                print(f" {highlight.bold('7')} - View Contact Sheet")
                valid_actions.append("7")
    print("     ------------")
    print(f" {highlight.bold('E')} - Edit Configs")
    valid_actions.append("E")
    print(f" {highlight.bold('Q')} - Quit")
    valid_actions.append("Q")
    print("----------------------------")
    return input("> ").upper()


def _facing(azimuth: str) -> str:
    if configs.get("FACING", "precision").lower() == "coarse":
        increment = 22.5
        directions = [
            "N",
            "NE",
            "E",
            "SE",
            "S",
            "SW",
            "W",
            "NW",
            "N",
        ]
    elif configs.get("FACING", "precision").lower() == "fine":
        increment = 11.25
        directions = [
            "N",
            "NNE",
            "NE",
            "ENE",
            "E",
            "ESE",
            "SE",
            "SSE",
            "S",
            "SSW",
            "SW",
            "WSW",
            "W",
            "WNW",
            "NW",
            "NNW",
            "N",
        ]
    try:
        azimuth = float(azimuth)
        if configs.get("FACING", "precision").lower() == "precise":
            bearing = f"{round(azimuth)}°"
        else:
            bearing = directions[math.ceil(math.floor((azimuth % 360) / increment) / 2)]
    except:
        bearing = ""
    return bearing


def _make_label(thephoto: dict) -> list:
    label = []
    line = []
    if thephoto["Subject"]:
        line.append(thephoto["Subject"])
    if thephoto["Description"]:
        line.append(thephoto["Description"])
    if line:
        label.append(": ".join(line))
    label.append(f"Original Photo: {thephoto['Photo']}")
    line = []
    if thephoto["Project"]:
        line.append(f"Project: {thephoto['Project']}")
    if thephoto["Site"]:
        line.append(f"Site: {thephoto['Site']}")
    if line:
        label.append("   ".join(line))
    line = []
    if thephoto["Facing"]:
        line.append(f"Facing {thephoto['Facing']}")
    if thephoto["GPS Coordinates"]:
        line.append(thephoto["GPS Coordinates"])
    if line:
        label.append(".  ".join(line))
    if thephoto["Timestamp"]:
        label.append(thephoto["Timestamp"])
    return label


def _read_configs() -> None:
    global configs
    if not Path("configs.ini").is_file():
        with open("configs.ini", "w") as f:
            pass
    configs.read("configs.ini")
    sections = ["EXIFTOOL", "DEFAULTS", "FACING", "RENAMING"]
    for each_section in sections:
        if not configs.has_section(each_section):
            configs.add_section(each_section)
    # EXIFTOOL section settings
    if not configs.has_option("EXIFTOOL", "exiftool"):
        if platform.system() == "Linux":
            configs.set("EXIFTOOL", "# Linux standard location of exiftool in $PATH")
            configs.set("EXIFTOOL", "exiftool", "exiftool")
        elif platform.system() == "Darwin":
            configs.set("EXIFTOOL", "# MacOS standard location of exiftool")
            configs.set("EXIFTOOL", "exiftool", "/usr/local/bin/exiftool")
        elif platform.system() == "Windows":
            configs.set(
                "EXIFTOOL",
                "# Windows recommended location of exiftool in PhotoCaptionTool folder",
            )
            configs.set("EXIFTOOL", "exiftool", "exiftool.exe")
    # DEFAULTS section settings
    if not configs.has_option("DEFAULTS", "papersize"):
        configs.set("DEFAULTS", "# papersize options are 'a4' and 'letter'")
        configs.set("DEFAULTS", "papersize", "a4")
    if not configs.has_option("DEFAULTS", "subjectdelimiter"):
        configs.set("DEFAULTS", "subjectdelimiter", ":")
    if not configs.has_option("DEFAULTS", "photographer"):
        configs.set("DEFAULTS", "photographer", "")
    if not configs.has_option("DEFAULTS", "project"):
        configs.set("DEFAULTS", "project", "")
    if not configs.has_option("DEFAULTS", "site"):
        configs.set("DEFAULTS", "site", "")
    # FACING section settings
    if not configs.has_option("FACING", "precision"):
        configs.set("FACING", "# precision options are")
        configs.set("FACING", "#   'coarse' (N, NE, E, SE, S, SW, W, NW)")
        configs.set("FACING", "#   'fine' (N, NNE, NE, ENE, E, and so-on)")
        configs.set("FACING", "#   'precise' (the actual bearing, in degrees)")
        configs.set("FACING", "precision", "coarse")
    # RENAMING section settings
    if not configs.has_option("RENAMING", "format"):
        configs.set("RENAMING", "# format options are")
        configs.set("RENAMING", "#   '1' (Subject -- Photographer_Photo.jpg)")
        configs.set("RENAMING", "#   '2' (Site_Subject_Sequence.jpg)")
        configs.set("RENAMING", "format", "1")
    with open("configs.ini", "w") as f:
        configs.write(f)
    if not shutil.which(configs.get("EXIFTOOL", "exiftool")):
        print("\n")
        print(
            f"{highlight.red('NOTICE:')}\nExifTool not found at the location indicated in the configs.ini file.\nInstall it according to the instructions in the README file or enter “E” to edit the configs and update its path."
        )


def _replace_invalid_filename_characters(thestring: str) -> str:
    invalid_chars = ["<", ">", ":", '"', "/", "\\", "|", "?", "*"]
    for each_char in invalid_chars:
        thestring = thestring.replace(each_char, " ")
    return thestring


def annotate_photos() -> None:
    global all_images_exif_data
    csv_file = Path(images_directory) / "Photo Log.csv"
    if not csv_file.is_file():
        main()
    output_dir = Path(images_directory) / "Annotated Photos"
    if output_dir.is_dir():
        if (
            input(
                f"“{output_dir}” already exists. Type “Y” to overwrite it.\n> "
            ).upper()
            != "Y"
        ):
            main()
        else:
            shutil.rmtree(output_dir)
    output_dir.mkdir(exist_ok=True)
    with open(csv_file, "r") as f:
        reader = csv.DictReader(f)
        i = 1
        for each_photo in reader:
            print(f"{i}: Annotating photo {each_photo['Photo']}.")
            label = _make_label(each_photo)
            orientation = all_images_exif_data[each_photo["Photo"]]["orientation"]
            if not orientation:
                orientation = "1"
            img = Image.open(Path(images_directory) / each_photo["Photo"])
            img = img.rotate(rotation.index(orientation) * 90, expand=True)
            img = ImageOps.pad(img, (img.width, img.height + 320), centering=(0, 0))
            img_annotation = ImageDraw.Draw(img, mode="RGB")
            try:
                thefont = ImageFont.truetype("Helvetica.ttc", 46)
            except:
                thefont = ImageFont.truetype("arial.ttf", 46)
            img_annotation.text(
                (20, img.height - 290),
                "\n".join(label),
                font=thefont,
                fill=(255, 255, 255),
                spacing=20,
            )
            filename = ".".join(each_photo["Photo"].split(".")[:-1])
            if configs.get("RENAMING", "format") == "1":
                if each_photo["Subject"]:
                    if each_photo["Photographer"]:
                        filename = f"{each_photo['Subject']} -- {each_photo['Photographer']}_{filename}"
                    else:
                        filename = f"{each_photo['Subject']} -- {filename}"
            elif configs.get("RENAMING", "format") == "2":
                filename = f"{each_photo['Site']}_{each_photo['Subject']}_{each_photo['Sequence']}.jpg"
            filename = f"{filename}_Annotated.jpg"
            img.save(
                Path(output_dir) / filename,
                quality=80,
                optimize=True,
                progressive=True,
            )
            img.close()
            caption = _build_new_caption(
                each_photo["Project"],
                each_photo["Site"],
                each_photo["Subject"],
                each_photo["Description"],
            )
            _ = subprocess.run(
                [
                    Path(configs.get("EXIFTOOL", "exiftool")),
                    "-tagsFromFile",
                    Path(images_directory) / each_photo["Photo"],
                    "-all:all",
                    Path(output_dir) / filename,
                    f"-artist={each_photo['Photographer']}",
                    f"-imagedescription={caption}",
                    f"-caption-abstract={caption}",
                    f"-description={caption}",
                    "--usercomment",
                    "-overwrite_original",
                    Path(output_dir) / filename,
                    "-overwrite_original",
                ],
                capture_output=True,
            )
            i += 1
    main()


def create_csv() -> None:
    headers = [
        "Photo",
        "Photographer",
        "Project",
        "Site",
        "Timestamp",
        "GPS Coordinates",
        "Facing",
        "Subject",
        "Description",
        "Sequence",
    ]
    if not all_images_exif_data:
        main()
    csv_file = Path(images_directory) / "Photo Log.csv"
    if (
        csv_file.is_file()
        and input(f"“{csv_file}” already exists. Type “Y” to overwrite it.\n> ").upper()
        != "Y"
    ):
        main()
    data_for_csv = []
    errors = []
    sequence = 1
    for each_image in all_images_exif_data:
        image_data = {"Photo": each_image}
        try:
            image_data["Photographer"] = configs.get("DEFAULTS", "photographer")
            image_data["Project"] = configs.get("DEFAULTS", "project")
            image_data["Site"] = configs.get("DEFAULTS", "site")
            if not configs.get("DEFAULTS", "photographer"):
                photographer = []
                if all_images_exif_data[each_image]["artist"]:
                    photographer.append(all_images_exif_data[each_image]["artist"])
                if (
                    all_images_exif_data[each_image]["creator"]
                    and all_images_exif_data[each_image]["creator"] != photographer[0]
                ):
                    photographer.append(all_images_exif_data[each_image]["creator"])
                image_data["Photographer"] = ", ".join(photographer)
            thedate = ""
            thedate = (
                all_images_exif_data[each_image]["datetimeoriginal"]
                .split(" ")[0]
                .replace(":", "-")
            )
            thetime = ""
            thetime = all_images_exif_data[each_image]["datetimeoriginal"].split(" ")[1]
            image_data["Timestamp"] = f"{thedate} {thetime}"
            image_data["GPS Coordinates"] = all_images_exif_data[each_image][
                "gpsposition"
            ]
            image_data["Facing"] = _facing(
                all_images_exif_data[each_image]["gpsimgdirection"]
            )
            caption = ""
            # Photo taken with iOS Camera.app:
            if all_images_exif_data[each_image]["imagedescription"]:
                caption = all_images_exif_data[each_image]["imagedescription"]
            # Photo taken with Theodolite.app:
            if all_images_exif_data[each_image]["usercomment"]:
                caption = all_images_exif_data[each_image]["usercomment"]
            image_data["Subject"] = ""
            image_data["Description"] = ""
            if caption.find(configs.get("DEFAULTS", "subjectdelimiter")) > 1:
                image_data["Subject"] = _replace_invalid_filename_characters(
                    caption.split(configs.get("DEFAULTS", "subjectdelimiter"))[
                        0
                    ].strip()
                )
                image_data["Description"] = (
                    configs.get("DEFAULTS", "subjectdelimiter")
                    .join(
                        caption.split(configs.get("DEFAULTS", "subjectdelimiter"))[1:]
                    )
                    .strip()
                )
            else:
                image_data["Description"] = caption
            image_data["Sequence"] = f"{sequence:03d}"
        except:
            errors.append(each_image)
        data_for_csv.append(image_data)
        sequence += 1
    with open(csv_file, "w", newline="") as f:
        csv_out = csv.DictWriter(f, fieldnames=headers)
        csv_out.writeheader()
        csv_out.writerows(data_for_csv)
    print("“Photo Log.csv” created.")
    if errors:
        print("\n")
        print(
            f"{highlight.red('NOTICE:')} The following photos had corrupted or incomplete EXIF data."
        )
        for each_error in errors:
            print(f"- {each_error}")
        print("\n")
    main()


def create_word_doc() -> None:
    csv_file = Path(images_directory) / "Photo Log.csv"
    if not csv_file.is_file():
        main()
    word_doc = Path(images_directory) / "Contact Sheet.docx"
    if (
        word_doc.is_file()
        and input(f"“{word_doc}” already exists. Type “Y” to overwrite it.\n> ").upper()
        != "Y"
    ):
        main()
    # Create a new Word document on A4 paper
    document = Document()
    section = document.sections[0]
    # Default papersize is a4
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(12)
    section.right_margin = Mm(12)
    section.top_margin = Mm(12)
    section.bottom_margin = Mm(12)
    photowidth = Mm(90)
    if configs.get("DEFAULTS", "papersize").lower() == "letter":
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        photowidth = Inches(3.5)
    with open(csv_file, "r") as f:
        reader = csv.DictReader(f)
        i = 1
        for each_photo in reader:
            print(f"{i}: Adding photo {each_photo['Photo']} to contact sheet.")
            document.add_picture(
                str(Path(images_directory) / each_photo["Photo"]),
                width=photowidth,
            )
            label = _make_label(each_photo)
            document.add_paragraph("\n".join(label))
            document.add_paragraph()
            if not i % 2:
                document.add_page_break()
            i += 1
    document.save(word_doc)
    print("“Contact Sheet.docx” created.")
    main()


def edit_configs() -> None:
    # exiftool
    print("Enter the path to your exiftool installation.")
    exiftool = input(f"[{configs.get('EXIFTOOL', 'exiftool')}] > ")
    if exiftool:
        configs.set("EXIFTOOL", "exiftool", exiftool)

    # papersize
    print("Enter the paper size for Word docs.")
    print("(options are a4 or letter)")
    papersize = input(f"[{configs.get('DEFAULTS', 'papersize')}] > ").lower()
    if papersize:
        while papersize not in ["a4", "letter"]:
            print("Invalid option entered. Please enter either a4 or letter.")
            papersize = input(f"[{configs.get('DEFAULTS', 'papersize')}] > ").lower()
            if not papersize:
                papersize = configs.get("DEFAULTS", "papersize")
        configs.set("DEFAULTS", "papersize", papersize)

    # subjectdelimiter
    print("Enter the delimiter between subject and description in your photo captions.")
    subjectdelimiter = input(f"[{configs.get('DEFAULTS', 'subjectdelimiter')}] > ")
    if subjectdelimiter:
        configs.set("DEFAULTS", "subjectdelimiter", subjectdelimiter)

    # photographer
    print(
        "Enter the name or initials of the photographer. Dash (-) to clear this setting."
    )
    photographer = input(f"[{configs.get('DEFAULTS', 'photographer')}] > ")
    if photographer:
        if photographer == "-":
            photographer = ""
        configs.set("DEFAULTS", "photographer", photographer)

    # project
    print("Enter the name of of the project. Dash (-) to clear this setting.")
    project = input(f"[{configs.get('DEFAULTS', 'project')}] > ")
    if project:
        if project == "-":
            project = ""
        configs.set("DEFAULTS", "project", project)

    # site
    print("Enter the name of of the site. Dash (-) to clear this setting")
    site = input(f"[{configs.get('DEFAULTS', 'site')}] > ")
    if site:
        if site == "-":
            site = ""
        configs.set("DEFAULTS", "site", site)

    # precision
    print("Enter the level of precision for the direction the photographs are facing.")
    print("(options are coarse, fine, or precise)")
    precision = input(f"[{configs.get('FACING', 'precision')}] > ").lower()
    if precision:
        while precision not in ["coarse", "fine", "precise"]:
            print(
                "Invalid option entered. Please enter either coarse, fine, or precise."
            )
            precision = input(f"[{configs.get('FACING', 'precision')}] > ").lower()
            if not precision:
                precision = configs.get("FACING", "precision")
        configs.set("FACING", "precision", precision)

    # format
    print("Enter the format for renamed photos.")
    print("(options are 1 or 2)")
    format = input(f"[{configs.get('RENAMING', 'format')}] > ")
    if format:
        while format not in ["1", "2"]:
            print("Invalid option entered. Please enter either 1 or 2.")
            format = input(f"[{configs.get('RENAMING', 'format')}] > ")
            if not format:
                format = configs.get("RENAMING", "format")
        configs.set("RENAMING", "format", format)

    with open("configs.ini", "w") as f:
        configs.write(f)
    main()


def load_photos() -> None:
    global images_directory
    global all_images_exif_data
    all_images_exif_data = {}
    images_directory = input(
        "Enter the Images Folder (type the path or drag the folder onto here)\n> "
    )
    if not images_directory:
        main()
    if images_directory[0] == "'" and images_directory[-1] == "'":
        images_directory = images_directory.strip("'")
    elif images_directory[0] == '"' and images_directory[-1] == '"':
        images_directory = images_directory.strip('"')
    else:
        images_directory = images_directory.strip()
    if "PosixPath" in str(type(Path())):
        images_directory = images_directory.replace("\\", "")
    if images_directory.startswith("~"):
        images_directory = images_directory.replace("~", str(Path.home()), 1)
    images = []
    items = Path(images_directory).glob("*.*")
    for each_item in items:
        if each_item.name[0] != "." and re.match(
            r"(.*\.jpe?g)|(.*\.heic)", each_item.name, flags=re.IGNORECASE
        ):
            images.append(each_item.name)
            images.sort()
    if len(images) == 0:
        images_directory = ""
        print("\n")
        print(
            f"{highlight.red('NOTICE:')} There were no valid JPEG images in the selected directory."
        )
    else:
        bad_photos = []
        tags = [
            "-datetimeoriginal",
            "-artist",
            "-creator",
            "-imagedescription",
            "-usercomment",
            "-gpsposition",
            "-gpsimgdirection",
            "-orientation#",
        ]
        for each_image in images:
            exif_data = subprocess.run(
                [
                    Path(configs.get("EXIFTOOL", "exiftool")),
                    "-T",
                    *tags,
                    Path(images_directory) / each_image,
                ],
                capture_output=True,
                text=True,
            )
            if exif_data.returncode == 0:
                exif_data = exif_data.stdout.strip().split("\t")
                all_images_exif_data[each_image] = dict(
                    zip(
                        [x.strip("-").strip("#") for x in tags],
                        ["" if x == "-" else x for x in exif_data],
                    )
                )
            else:
                # Windows will cause exiftool to choke on unicode characters in the file name.
                bad_photos.append(each_image)
        if bad_photos:
            print("\n")
            print(
                f"{highlight.red('NOTICE:')} The following photos couldn’t be read. They probably have unicode characters in their file names."
            )
            for each_bad_photo in bad_photos:
                print(f" - {each_bad_photo}")
    main()


def rename_photos() -> None:
    csv_file = Path(images_directory) / "Photo Log.csv"
    if not csv_file.is_file():
        main()
    output_dir = Path(images_directory) / "Renamed Photos"
    if output_dir.is_dir():
        if (
            input(
                f"“{output_dir}” already exists. Type “Y” to overwrite it.\n> "
            ).upper()
            != "Y"
        ):
            main()
        else:
            shutil.rmtree(output_dir)
    output_dir.mkdir(exist_ok=True)
    with open(csv_file, "r") as f:
        reader = csv.DictReader(f)
        i = 1
        for each_photo in reader:
            print(f"{i}: Renaming photo {each_photo['Photo']}.")
            filename = ".".join(each_photo["Photo"].split(".")[:-1])
            if configs.get("RENAMING", "format") == "1":
                if each_photo["Subject"]:
                    if each_photo["Photographer"]:
                        filename = f"{each_photo['Subject']} -- {each_photo['Photographer']}_{filename}"
                    else:
                        filename = f"{each_photo['Subject']} -- {filename}"
            elif configs.get("RENAMING", "format") == "2":
                filename = f"{each_photo['Site']}_{each_photo['Subject']}_{each_photo['Sequence']}.jpg"
            filename = f"{filename}.jpg"
            if each_photo["Photo"].split(".")[-1].upper() == "HEIC":
                img = Image.open(Path(images_directory) / each_photo["Photo"])
                img.save(
                    Path(output_dir) / filename,
                    quality=80,
                    optimize=True,
                    progressive=True,
                )
                img.close()
            else:
                shutil.copy2(
                    Path(images_directory) / each_photo["Photo"],
                    Path(output_dir) / filename,
                )
            caption = _build_new_caption(
                each_photo["Project"],
                each_photo["Site"],
                each_photo["Subject"],
                each_photo["Description"],
            )
            _ = subprocess.run(
                [
                    Path(configs.get("EXIFTOOL", "exiftool")),
                    f"-artist={each_photo['Photographer']}",
                    f"-imagedescription={caption}",
                    f"-caption-abstract={caption}",
                    f"-description={caption}",
                    "--usercomment",
                    "-overwrite_original",
                    Path(output_dir) / filename,
                ],
                capture_output=True,
            )
            i += 1
    main()


def view_csv_file() -> None:
    thefile = Path(images_directory) / "Photo Log.csv"
    try:
        os.startfile(thefile)
    except:
        os.system(f'open "{thefile}"')
    main()


def view_word_doc() -> None:
    thefile = Path(images_directory) / "Contact Sheet.docx"
    try:
        os.startfile(thefile)
    except:
        os.system(f'open "{thefile}"')
    main()


def main() -> None:
    if platform.system() == "Windows":
        os.system("highlight")
    _read_configs()
    action = ""
    while action not in valid_actions:
        action = _display_menu()
    if action == "1":
        load_photos()
    elif action == "2":
        create_csv()
    elif action == "3":
        view_csv_file()
    elif action == "4":
        rename_photos()
    elif action == "5":
        annotate_photos()
    elif action == "6":
        create_word_doc()
    elif action == "7":
        view_word_doc()
    elif action == "E":
        edit_configs()
    elif action == "Q":
        exit()


if __name__ == "__main__":
    main()
